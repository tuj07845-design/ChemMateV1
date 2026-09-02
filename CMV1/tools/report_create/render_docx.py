from __future__ import annotations

import os
from datetime import datetime
from pathlib import Path
from typing import Any

CN_BODY_FONT = "宋体"          # Word 正文中文
CN_HEAD_FONT = "黑体"          # Word 标题中文
LATIN_BODY_FONT = "Times New Roman"
LATIN_HEAD_FONT = "Arial"

# ============================================================
# Word 渲染
# ============================================================



# glm —— 段落样式级中西文字体（w:rFonts/@w:eastAsia）
def _docx_style_font(style, east: str = CN_BODY_FONT, latin: str = LATIN_BODY_FONT) -> None:
    """给段落样式设置中西文字体（w:rFonts/@w:eastAsia 决定中文字体）。"""
    from docx.oxml.ns import qn
    style.font.name = latin  # 先设西文，确保 rFonts 存在
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east)


# glm —— run 级中西文字体/字号/加粗/颜色
def _docx_run_font(run, size=None, bold=None, color=None,
                   east: str = CN_BODY_FONT, latin: str = LATIN_BODY_FONT) -> None:
    """给 run 设置中西文字体、字号、加粗、颜色。"""
    from docx.oxml.ns import qn
    run.font.name = latin
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


# glm —— 首行缩进按字符数（firstLineChars 为主，中文两字符习惯）
def _docx_first_line_indent(paragraph, chars: int = 2) -> None:
    """首行缩进按字符数：firstLineChars 为准（中文 Word 习惯），firstLine 兜底。"""
    from docx.oxml.ns import qn
    ind = paragraph._p.get_or_add_pPr().get_or_add_ind()
    ind.set(qn("w:firstLine"), str(chars * 240))       # 12pt 字号时一字符 240 twips
    ind.set(qn("w:firstLineChars"), str(chars * 100))


# glm —— 页脚居中“第 X 页”（PAGE 域）
def _docx_add_page_number(footer_paragraph) -> None:
    """页脚居中插入“第 X 页”（PAGE 域）。"""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _docx_run_font(footer_paragraph.add_run("第 "), size=Pt(9))
    r = footer_paragraph.add_run()
    fld_b = OxmlElement("w:fldChar")
    fld_b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_e = OxmlElement("w:fldChar")
    fld_e.set(qn("w:fldCharType"), "end")
    r._r.append(fld_b)
    r._r.append(instr)
    r._r.append(fld_e)
    r.font.size = Pt(9)
    _docx_run_font(footer_paragraph.add_run(" 页"), size=Pt(9))


def _render_docx(title: str, sections: list, output_path: Path, author: str) -> Path:
    """渲染成 .docx（排版约定见模块文档：宋体正文/黑体标题/1.5 倍行距/首行缩进）。"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement

    doc = Document()

    # glm —— 页面：A4，上下 2.54cm，左右 3.18cm（Word 常规页边距）
    sec0 = doc.sections[0]
    sec0.page_width, sec0.page_height = Cm(21.0), Cm(29.7)
    sec0.top_margin = sec0.bottom_margin = Cm(2.54)
    sec0.left_margin = sec0.right_margin = Cm(3.18)

    # glm —— 样式：正文宋体小四、1.5 倍行距；标题黑体加粗；项目符号宋体
    normal = doc.styles["Normal"]
    _docx_style_font(normal, CN_BODY_FONT, LATIN_BODY_FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size, before, after in (
        ("Title", 22, 0, 12),
        ("Heading 1", 16, 14, 8),
        ("Heading 2", 14, 10, 6),
        ("Heading 3", 12, 8, 4),
    ):
        st = doc.styles[style_name]
        _docx_style_font(st, CN_HEAD_FONT, LATIN_HEAD_FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    list_bullet = doc.styles["List Bullet"]
    _docx_style_font(list_bullet, CN_BODY_FONT, LATIN_BODY_FONT)
    list_bullet.font.size = Pt(12)
    list_bullet.paragraph_format.line_spacing = 1.5
    list_bullet.paragraph_format.space_after = Pt(0)

    # 核心属性
    try:
        doc.core_properties.author = author
        doc.core_properties.title = title
    except Exception:
        pass

    # glm —— 封面：标题居中（Title 样式）
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # glm —— 副信息：生成时间 + 作者（小字灰色居中）
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _docx_run_font(
        meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}    {author}"),
        size=Pt(10), color=RGBColor(0x66, 0x66, 0x66),
    )

    # glm —— 页脚页码
    _docx_add_page_number(sec0.footer.paragraphs[0])

    for sec in sections:
        stype = sec.get("type")

        if stype == "heading":
            level = int(sec.get("level", 1) or 1)
            level = max(1, min(3, level))
            doc.add_heading(str(sec.get("text", "")), level=level)

        elif stype == "paragraph":
            # glm —— 正文：两端对齐 + 首行缩进两字符
            p = doc.add_paragraph(str(sec.get("text", "")))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _docx_first_line_indent(p, 2)

        elif stype == "bullets":
            items = sec.get("items") or []
            if isinstance(items, list):
                for it in items:
                    bp = doc.add_paragraph(str(it), style="List Bullet")
                    bp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        elif stype == "table":
            headers = sec.get("headers") or []
            rows = sec.get("rows") or []
            if not isinstance(headers, list) or not isinstance(rows, list):
                continue
            n_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
            if n_cols == 0:
                continue
            table = doc.add_table(rows=1 + len(rows), cols=n_cols)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # glm —— 表格排版：宋体五号，表头加粗居中；行禁止跨页断开
            for i, row in enumerate([headers] + list(rows)):
                vals = row if isinstance(row, list) else []
                trpr = table.rows[i]._tr.get_or_add_trPr()
                trpr.append(OxmlElement("w:cantSplit"))
                for j in range(n_cols):
                    val = vals[j] if j < len(vals) else ""
                    cell = table.rows[i].cells[j]
                    cell.text = "" if val is None else str(val)
                    for p in cell.paragraphs:
                        p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if i == 0
                                       else WD_ALIGN_PARAGRAPH.LEFT)
                        p.paragraph_format.line_spacing = 1.0
                        for r in p.runs:
                            _docx_run_font(r, size=Pt(10.5), bold=(i == 0))
            # glm —— 表头跨页重复
            table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))

        elif stype == "image":
            img_path = sec.get("path", "")
            caption = sec.get("caption", "")
            if not img_path or not os.path.isfile(img_path):
                # 图片缺失不中断，插一行提示
                p = doc.add_paragraph()
                _docx_run_font(p.add_run(f"[图片缺失：{img_path}]"),
                               color=RGBColor(0xCC, 0x00, 0x00))
            else:
                try:
                    doc.add_picture(img_path, width=Inches(5.5))
                except Exception:
                    doc.add_picture(img_path)
                # glm —— 图注：小字灰色居中
                if caption:
                    cap = doc.add_paragraph(str(caption))
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in cap.runs:
                        _docx_run_font(r, size=Pt(9), color=RGBColor(0x66, 0x66, 0x66))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path

